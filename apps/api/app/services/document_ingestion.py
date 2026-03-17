from __future__ import annotations

import io
import zipfile
from dataclasses import dataclass
from pathlib import Path
from uuid import uuid4

import fitz
from docx import Document

from app.schemas.domain import TextFragment
from app.utils.text import clean_text, split_lines

SUPPORTED_FILE_TYPES = {"pdf", "docx", "txt"}


class DocumentExtractionError(ValueError):
    pass


@dataclass
class ExtractedDocument:
    filename: str
    file_type: str
    raw_text: str
    fragments: list[TextFragment]


class DocumentExtractor:
    def __init__(self, max_size_bytes: int) -> None:
        self.max_size_bytes = max_size_bytes

    def extract(self, filename: str, payload: bytes) -> ExtractedDocument:
        safe_name = Path(filename).name
        extension = Path(safe_name).suffix.lower().lstrip(".")
        if extension not in SUPPORTED_FILE_TYPES:
            raise DocumentExtractionError(f"Unsupported file type: {extension}")
        if len(payload) > self.max_size_bytes:
            raise DocumentExtractionError("Uploaded file exceeds the configured size limit")

        if extension == "pdf":
            self._validate_pdf(payload)
            fragments = self._extract_pdf(payload)
        elif extension == "docx":
            self._validate_docx(payload)
            fragments = self._extract_docx(payload)
        else:
            fragments = self._extract_txt(payload)

        raw_text = "\n\n".join(fragment.text for fragment in fragments if fragment.text.strip())
        if not fragments or not clean_text(raw_text):
            raise DocumentExtractionError("No readable text found in the uploaded document")
        return ExtractedDocument(
            filename=safe_name,
            file_type=extension,
            raw_text=clean_text(raw_text),
            fragments=fragments,
        )

    def _validate_pdf(self, payload: bytes) -> None:
        if not payload.startswith(b"%PDF"):
            raise DocumentExtractionError("The uploaded file does not look like a valid PDF")

    def _validate_docx(self, payload: bytes) -> None:
        if not zipfile.is_zipfile(io.BytesIO(payload)):
            raise DocumentExtractionError("The uploaded file does not look like a valid DOCX")

    def _extract_pdf(self, payload: bytes) -> list[TextFragment]:
        document = fitz.open(stream=payload, filetype="pdf")
        fragments: list[TextFragment] = []
        order_index = 0
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            blocks = page.get_text("blocks")
            for block_index, block in enumerate(blocks):
                text = clean_text(block[4] or "")
                if not text:
                    continue
                fragments.append(
                    TextFragment(
                        id=str(uuid4()),
                        section_name="unclassified",
                        order_index=order_index,
                        text=text,
                        source_type="pdf",
                        page_number=page_index + 1,
                        block_index=block_index,
                    )
                )
                order_index += 1
        return fragments

    def _extract_docx(self, payload: bytes) -> list[TextFragment]:
        document = Document(io.BytesIO(payload))
        fragments: list[TextFragment] = []
        order_index = 0
        for paragraph in document.paragraphs:
            text = clean_text(paragraph.text)
            if not text:
                continue
            section_name = "heading" if paragraph.style and "heading" in paragraph.style.name.lower() else "unclassified"
            fragments.append(
                TextFragment(
                    id=str(uuid4()),
                    section_name=section_name,
                    order_index=order_index,
                    text=text,
                    source_type="docx",
                    block_index=order_index,
                )
            )
            order_index += 1
        return fragments

    def _extract_txt(self, payload: bytes) -> list[TextFragment]:
        text = payload.decode("utf-8", errors="ignore")
        paragraphs = [clean_text(chunk) for chunk in text.replace("\r", "").split("\n\n")]
        fragments: list[TextFragment] = []
        for index, paragraph in enumerate(paragraphs):
            if not paragraph:
                continue
            lines = split_lines(paragraph)
            section_name = "heading" if len(lines) == 1 and len(lines[0]) < 40 else "unclassified"
            fragments.append(
                TextFragment(
                    id=str(uuid4()),
                    section_name=section_name,
                    order_index=index,
                    text=paragraph,
                    source_type="txt",
                    block_index=index,
                )
            )
        return fragments
